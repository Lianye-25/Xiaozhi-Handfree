"""生成 HANDSFREE 五维中枢使用说明 Word 文档。"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_PATH = r"./shuoming.docx"


def set_cell_shading(cell, color):
    """设置单元格底色。"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_table(doc, headers, rows, col_widths=None):
    """添加表格。"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "4472C4")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    """添加标题。"""
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, size=10):
    """添加段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def main():
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 封面标题 =====
    title = doc.add_heading('HANDSFREE 五维中枢', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('语音指令使用说明（v2.0）')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()

    # ===== 联系人名单 =====
    add_heading(doc, '联系人名单', level=1)
    add_para(doc, '以下是当前系统中已配置的联系人，语音指令中的人名需与列表匹配：')

    add_table(doc,
        ['类型', '姓名', '身份说明', '通知方式'],
        [
            ['上级', '张总', '项目负责人', '邮件 (zhang@qq.com)'],
            ['上级', '李经理', '部门经理', '邮件 (li@qq.com)'],
            ['上级', '王总监', '技术总监', '邮件 (wang@company.com)'],
            ['下属', '小王', '开发工程师', '微信'],
            ['下属', '小明', '数据分析师', '邮件 (xiaoming@qq.com)'],
            ['下属', '小红', 'UI设计师', '微信'],
            ['下属', '小莲', '—', '邮件 (xiaolian@example.com)'],
            ['Agent', '代码审查Agent', '代码审查', 'LLM 自动执行'],
            ['Agent', '文档润色Agent', '文档润色', 'LLM 自动执行'],
            ['Agent', '（动态创建）', '用户说什么就创建什么', 'LLM 自动执行'],
        ],
        col_widths=[2, 3, 3, 5]
    )

    # ===== 系统概述 =====
    add_heading(doc, '系统概述', level=1)
    add_para(doc, 'HANDSFREE 五维中枢将语音指令分为五个方向：')
    add_para(doc, '↑ UP（上报）—— 向上级汇报，自动发邮件')
    add_para(doc, '↓ DOWN（委派）—— 向下属分派任务，自动发邮件')
    add_para(doc, '← LEFT（存档）—— 归档已完成任务')
    add_para(doc, '→ RIGHT（Agent分派）—— AI Agent 后台异步执行任务')
    add_para(doc, '中 CENTER（中枢）—— 信息查询，无特定方向')

    doc.add_paragraph()
    add_para(doc, '完整工作流示例：', bold=True)
    add_para(doc, '1. 你对着小智说："交给代码审查Agent检查这段代码"')
    add_para(doc, '2. 系统解析为 RIGHT 方向，在后台启动 Agent 执行任务')
    add_para(doc, '3. Agent 调用 DeepSeek API 完成代码审查（约5-30秒）')
    add_para(doc, '4. 小智主动播报："代码审查Agent已完成，结果为：……"')
    add_para(doc, '5. 你听后决定："存档这个结果" → 任务归档完成')
    add_para(doc, '6. 或者："汇报给张总审查结果" → 邮件发给张总')
    add_para(doc, '7. 或者："分派给小莲根据结果修改代码" → 邮件发给小莲')

    # ===== UP 上报 =====
    add_heading(doc, '↑ UP — 向上汇报', level=1)
    add_para(doc, '干什么用：把消息、报告、进展同步给上级。', bold=True)
    add_para(doc, '触发关键词：上报、汇报、提交、报告、告诉、通知、呈报、报给、发给、反馈给')
    add_para(doc, '执行动作：自动发送汇报邮件给匹配的上级。')

    add_table(doc,
        ['场景', '语音指令', '解析结果'],
        [
            ['汇报进展', '汇报给张总今天的项目进展已完成', 'UP → 张总 → 邮件发送'],
            ['提交周报', '报告给李经理周报已写好', 'UP → 李经理 → 邮件发送'],
            ['提交方案', '提交给王总监架构方案请审批', 'UP → 王总监 → 邮件发送'],
            ['通知变更', '通知张总会议推迟到下午三点', 'UP → 张总 → 邮件发送'],
            ['口语汇报', '报给张总项目已上线', 'UP → 张总 → 邮件发送'],
            ['反馈问题', '反馈给李经理客户有新需求', 'UP → 李经理 → 邮件发送'],
        ],
        col_widths=[2, 5, 5]
    )
    add_para(doc, '说明：说人名时要尽量说全称（"张总"而非"张"），系统会自动匹配联系人。对方邮箱地址需在 contacts.json 中配置。')

    # ===== DOWN 委派 =====
    add_heading(doc, '↓ DOWN — 向下委派', level=1)
    add_para(doc, '干什么用：把任务分派给下属去执行。', bold=True)
    add_para(doc, '触发关键词：分派、委派、分配、交给、让、安排、派给、下发、布置、叫')
    add_para(doc, '执行动作：自动发送任务委派邮件给匹配的下属。')

    add_table(doc,
        ['场景', '语音指令', '解析结果'],
        [
            ['分派任务', '分派给小王整理会议纪要', 'DOWN → 小王 → 邮件发送'],
            ['委派任务', '委派给小明做数据分析报告', 'DOWN → 小明 → 邮件发送'],
            ['安排工作', '安排小红设计新的登录页面', 'DOWN → 小红 → 仅微信，邮件跳过'],
            ['口语指派', '让小王去测试新功能', 'DOWN → 小王 → 邮件发送'],
            ['委派小莲', '分派给小莲整理会议纪要', 'DOWN → 小莲 → 邮件发送'],
            ['分配bug', '分配给小明修复这个bug', 'DOWN → 小明 → 邮件发送'],
        ],
        col_widths=[2, 5, 5]
    )
    add_para(doc, '说明：如果联系人只配置了微信、没配置邮箱（如小红），系统会优雅跳过邮件通道，不会报错。')

    # ===== LEFT 存档 =====
    add_heading(doc, '← LEFT — 任务存档', level=1)
    add_para(doc, '干什么用：将最近完成的任务归档（已完成 → 已归档），不涉及人际分发。', bold=True)
    add_para(doc, '触发关键词：存档、标记完成、标记任务完成、归档、关闭任务、结束任务、标记为完成')
    add_para(doc, '执行动作：找到最近一条 status=completed 的任务，将其标记为 archived。')

    add_table(doc,
        ['场景', '语音指令', '解析结果'],
        [
            ['存档结果', '存档这个结果', 'LEFT → 最近完成的任务已归档'],
            ['标记完成', '标记任务完成', 'LEFT → 最近完成的任务已归档'],
            ['归档文档', '归档这份文档', 'LEFT → 最近完成的任务已归档'],
            ['关闭任务', '关闭任务这个bug已修复', 'LEFT → 最近完成的任务已归档'],
            ['阶段完成', '标记为完成项目第一阶段', 'LEFT → 最近完成的任务已归档'],
        ],
        col_widths=[2, 5, 5]
    )
    add_para(doc, '说明：这个方向不需要指定联系人。如果没有待归档的已完成任务，系统会提示"没有待归档的任务"。')

    # ===== RIGHT Agent分派 =====
    add_heading(doc, '→ RIGHT — Agent 分派', level=1)
    add_para(doc, '干什么用：把任务交给 AI Agent 后台异步处理，完成后小智主动播报结果。', bold=True)
    add_para(doc, '触发关键词：Agent处理、AI处理、智能处理、机器处理、自动处理')
    add_para(doc, '执行动作：创建任务 → 启动独立子进程 → 调用 DeepSeek API → 完成后更新任务状态。')
    add_para(doc, '动态 Agent：在 contacts.json 预定义的 Agent 之外，用户可以随意说出新的 Agent 名称，系统自动创建。例如"翻译Agent""数据分析Agent"等，无需预先配置。', bold=True)

    add_table(doc,
        ['场景', '语音指令', '解析结果'],
        [
            ['审查代码', '交给代码审查Agent检查这段代码', 'RIGHT → 代码审查Agent → 后台执行'],
            ['润色文档', '让文档润色Agent处理这份报告', 'RIGHT → 文档润色Agent → 后台执行'],
            ['通用AI', 'AI处理这个数据表格', 'RIGHT → 通用Agent → 后台执行'],
            ['智能分析', '智能处理这份合同条款分析', 'RIGHT → 通用Agent → 后台执行'],
            ['动态翻译', '交给翻译Agent把这段话翻译成英文', 'RIGHT → 翻译Agent(动态) → 后台执行'],
            ['动态分析', '交给数据分析Agent分析这份报告', 'RIGHT → 数据分析Agent(动态) → 后台执行'],
        ],
        col_widths=[2, 5, 5]
    )
    add_para(doc, '异步回调流程：')
    add_para(doc, '1. 你说出指令 → 系统立即返回"Agent 任务已后台启动"')
    add_para(doc, '2. Agent 在后台独立进程执行（不阻塞你的操作）')
    add_para(doc, '3. 完成后自动更新任务状态为 completed')
    add_para(doc, '4. 小智机器人轮询发现任务完成 → TTS 语音播报结果')
    add_para(doc, '5. 你听完结果后决定下一步：存档 / 汇报给上级 / 委派给下属')

    # ===== CENTER 中枢 =====
    add_heading(doc, '中 CENTER — 中枢模式', level=1)
    add_para(doc, '干什么用：不包含任何方向词时默认进入中枢，通常是信息查询。', bold=True)
    add_para(doc, '执行动作：不触发任何分发或存档操作，仅返回解析结果。')

    add_table(doc,
        ['场景', '语音指令'],
        [
            ['查询任务', '今天有什么任务'],
            ['询问进度', '现在的进度怎么样了'],
            ['查看列表', '显示今天的任务列表'],
        ],
        col_widths=[3, 9]
    )

    # ===== 歧义处理 =====
    add_heading(doc, '歧义场景处理', level=1)
    add_para(doc, '如果一句话里包含多个方向关键词，系统会取第一个出现的方向作为主方向。', bold=True)

    add_table(doc,
        ['你说', '系统理解'],
        [
            ['汇报给张总然后分派给小王去执行', '主方向 = UP（汇报），"分派"被记录为歧义'],
            ['交给代码审查Agent优化', '方向 = RIGHT（因为目标是Agent）'],
            ['交给小王来做', '方向 = DOWN（因为目标是人）'],
            ['让文档润色Agent处理这份报告', '方向 = RIGHT（消歧规则：让+Agent）'],
        ],
        col_widths=[7, 7]
    )

    # ===== 任务状态查询 =====
    add_heading(doc, '任务状态查询命令', level=1)
    add_para(doc, '以下命令可在终端中运行，查看任务执行状态：')

    add_table(doc,
        ['命令', '用途'],
        [
            ['python scripts/task_status.py --pending', '查询所有待处理任务'],
            ['python scripts/task_status.py --running', '查询所有执行中任务'],
            ['python scripts/task_status.py --latest-completed', '查询最近完成的任务（小智轮询用）'],
            ['python scripts/task_status.py --list', '列出所有任务'],
            ['python scripts/task_status.py --archive <任务ID>', '手动归档指定任务'],
        ],
        col_widths=[8, 7]
    )
    add_para(doc, '任务状态流转：pending（待执行）→ running（执行中）→ completed（已完成）→ archived（已归档）')

    # ===== 配置说明 =====
    add_heading(doc, '配置说明', level=1)
    add_para(doc, '邮件发送配置（assets/config.json）：', bold=True)
    add_para(doc, '在 config.json 的 email 段填入 SMTP 信息（QQ邮箱需使用授权码而非登录密码）。')
    add_para(doc, '如果不配置 SMTP，系统仍可正常解析意图，只是跳过邮件发送步骤。')
    doc.add_paragraph()
    add_para(doc, 'LLM API 配置（assets/config.json）：', bold=True)
    add_para(doc, '在 config.json 的 llm 段填入 DeepSeek API Key 和模型名称。系统已默认指向 DeepSeek V4 Pro。')
    add_para(doc, '如果不配置 LLM API Key，Agent 分派（RIGHT）功能不可用。')

    # ===== 小贴士 =====
    add_heading(doc, '小贴士', level=1)
    add_para(doc, '1. 名字说清楚：系统通过联系人列表匹配，"张总"比"张"命中率更高。')
    add_para(doc, '2. 句式随意："汇报给张总……""给张总汇报……""和张总汇报……"都可以。')
    add_para(doc, '3. ASR 容错：如果语音识别把"张总"识别成"章总"，模糊匹配仍可能命中（置信度会降低）。')
    add_para(doc, '4. 没匹配到别担心：如果联系人不在名单里，系统会返回低置信度，不会报错。')
    add_para(doc, '5. 动态 Agent 万能：除了预定义的代码审查Agent和文档润色Agent，你说的任何 Agent 名称系统都会自动创建。')
    add_para(doc, '6. 邮件+Agent 双通道：汇报/委派走邮件，Agent分派走 LLM 后台执行，各司其职。')

    # 保存
    doc.save(OUT_PATH)
    print(f'文档已保存到: {OUT_PATH}')


if __name__ == '__main__':
    main()
